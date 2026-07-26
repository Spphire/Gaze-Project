from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "output" / "doc"
FIGURES = ROOT / "output" / "patent_figures"
OUTPUT_DOCX = OUT / "Gaze-WAM_发明专利技术交底书_代码对齐版.docx"

NAVY = "17202A"
INK = "25313B"
MUTED = "5C6873"
GREEN = "176B5B"
GREEN_LIGHT = "EAF5EF"
BLUE_LIGHT = "EAF2F8"
ORANGE_LIGHT = "FBF1E8"
PURPLE_LIGHT = "F3EEF8"
GRAY_LIGHT = "EEF1F3"
GRAY_MID = "D5DBDF"
WHITE = "FFFFFF"
RED_LIGHT = "FCEBEA"

CODE_COMMIT = "1e74d7a2c468801b8d6fae39d87c1fa1d55eef10"
SEARCH_CUTOFF = "2026-07-27"


def set_run_font(run, name: str = "Microsoft YaHei", size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = GRAY_MID, size: int = 6, outer: bool = True) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single" if outer or edge.startswith("inside") else "nil")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "nil")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=8.5)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_separate, value, fld_end])
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=8.5)


def add_hyperlink(paragraph, text: str, url: str, color: str = GREEN) -> None:
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color_el, underline])
    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_top_rule(paragraph, color: str = GREEN, size: int = 18) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), str(size))
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), color)
    p_bdr.append(top)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.22
    normal.paragraph_format.widow_control = True

    for style_name, size, color, before, after in (
        ("Title", 27, NAVY, 0, 14),
        ("Heading 1", 17, NAVY, 16, 8),
        ("Heading 2", 13.5, GREEN, 12, 5),
        ("Heading 3", 11.5, NAVY, 9, 3),
    ):
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name, base, size, color, italic in (
        ("Deck", "Normal", 12, MUTED, False),
        ("Caption Gaze", "Normal", 9, MUTED, False),
        ("Code Gaze", "Normal", 8.5, NAVY, False),
        ("Small Gaze", "Normal", 8, MUTED, False),
        ("Quote Gaze", "Normal", 10, NAVY, False),
    ):
        if style_name not in styles:
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = styles[base]
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.italic = italic
    styles["Caption Gaze"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["Caption Gaze"].paragraph_format.space_before = Pt(3)
    styles["Caption Gaze"].paragraph_format.space_after = Pt(9)
    styles["Code Gaze"].font.name = "Consolas"
    styles["Code Gaze"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Code Gaze"].paragraph_format.left_indent = Cm(0.5)
    styles["Code Gaze"].paragraph_format.right_indent = Cm(0.5)
    styles["Code Gaze"].paragraph_format.space_before = Pt(3)
    styles["Code Gaze"].paragraph_format.space_after = Pt(3)
    styles["Quote Gaze"].paragraph_format.left_indent = Cm(0.5)
    styles["Quote Gaze"].paragraph_format.right_indent = Cm(0.5)


def configure_sections(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.9)
        section.right_margin = Cm(1.7)
        section.header_distance = Cm(0.75)
        section.footer_distance = Cm(0.75)


def add_headers_footers(doc: Document) -> None:
    for section in doc.sections:
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        header = section.header
        p = header.paragraphs[0]
        p.clear()
        p.text = "GAZE-WAM  |  发明专利技术交底书  |  内部评审稿"
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        add_top_rule(p, color=GREEN, size=10)
        for run in p.runs:
            set_run_font(run, size=7.5, bold=True)
            run.font.color.rgb = RGBColor.from_string(MUTED)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.clear()
        add_page_number(fp)
        for run in fp.runs:
            run.font.color.rgb = RGBColor.from_string(MUTED)


def paragraph(container, text: str = "", *, bold_prefix: str | None = None, style: str | None = None):
    p = container.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix) :])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def bullet_list(container, items: Iterable[str]) -> None:
    for item in items:
        p = container.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        set_run_font(r)


def numbered_list(container, items: Iterable[str]) -> None:
    for item in items:
        p = container.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run_font(r)


def callout(container, title: str, body: str, fill: str = GREEN_LIGHT) -> None:
    table = container.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    remove_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 130, 180, 130, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    set_run_font(run, size=11, bold=True)
    run.font.color.rgb = RGBColor.from_string(GREEN if fill != RED_LIGHT else "A23B32")
    body_p = cell.add_paragraph()
    body_p.paragraph_format.space_after = Pt(0)
    run = body_p.add_run(body)
    set_run_font(run, size=9.5)


def add_table(
    container,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    widths_cm: Sequence[float] | None = None,
    font_size: float = 8.5,
):
    table = container.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = widths_cm is None
    set_table_borders(table)
    header = table.rows[0]
    repeat_table_header(header)
    prevent_row_split(header)
    for idx, text in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths_cm:
            cell.width = Cm(widths_cm[idx])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, size=font_size, bold=True)
        r.font.color.rgb = RGBColor.from_string(WHITE)
    for row_index, row_data in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        for idx, text in enumerate(row_data):
            cell = row.cells[idx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths_cm:
                cell.width = Cm(widths_cm[idx])
            if row_index % 2 == 1:
                set_cell_shading(cell, "F7F8F9")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(text))
            set_run_font(r, size=font_size)
    container.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_code_block(container, code: str) -> None:
    table = container.add_table(rows=1, cols=1)
    remove_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, GRAY_LIGHT)
    set_cell_margins(cell, 130, 160, 130, 160)
    p = cell.paragraphs[0]
    p.style = "Code Gaze"
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(code)
    set_run_font(run, "Consolas", 8.2)


def add_figure(doc: Document, filename: str, caption: str, width: float = 6.75) -> None:
    path = FIGURES / filename
    if not path.is_file():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cp = doc.add_paragraph(caption, style="Caption Gaze")
    cp.paragraph_format.keep_with_next = False


def m_run(text: str, plain: bool = False):
    run = OxmlElement("m:r")
    run_pr = OxmlElement("m:rPr")
    style = OxmlElement("m:sty")
    style.set(qn("m:val"), "p" if plain else "i")
    run_pr.append(style)
    run.append(run_pr)
    node = OxmlElement("m:t")
    node.text = text
    run.append(node)
    return run


def m_expr(value):
    if isinstance(value, str):
        return [m_run(value)]
    if isinstance(value, (list, tuple)):
        out = []
        for item in value:
            out.extend(m_expr(item))
        return out
    return [value]


def m_sub(base, subscript):
    node = OxmlElement("m:sSub")
    e = OxmlElement("m:e")
    sub = OxmlElement("m:sub")
    for item in m_expr(base):
        e.append(item)
    for item in m_expr(subscript):
        sub.append(item)
    node.extend([e, sub])
    return node


def m_sup(base, superscript):
    node = OxmlElement("m:sSup")
    e = OxmlElement("m:e")
    sup = OxmlElement("m:sup")
    for item in m_expr(base):
        e.append(item)
    for item in m_expr(superscript):
        sup.append(item)
    node.extend([e, sup])
    return node


def m_subsup(base, subscript, superscript):
    node = OxmlElement("m:sSubSup")
    e = OxmlElement("m:e")
    sub = OxmlElement("m:sub")
    sup = OxmlElement("m:sup")
    for item in m_expr(base):
        e.append(item)
    for item in m_expr(subscript):
        sub.append(item)
    for item in m_expr(superscript):
        sup.append(item)
    node.extend([e, sub, sup])
    return node


def m_frac(numerator, denominator):
    node = OxmlElement("m:f")
    num = OxmlElement("m:num")
    den = OxmlElement("m:den")
    for item in m_expr(numerator):
        num.append(item)
    for item in m_expr(denominator):
        den.append(item)
    node.extend([num, den])
    return node


def m_sqrt(expression):
    node = OxmlElement("m:rad")
    props = OxmlElement("m:radPr")
    hide = OxmlElement("m:degHide")
    hide.set(qn("m:val"), "1")
    props.append(hide)
    degree = OxmlElement("m:deg")
    e = OxmlElement("m:e")
    for item in m_expr(expression):
        e.append(item)
    node.extend([props, degree, e])
    return node


def m_delim(expression, left: str = "(", right: str = ")"):
    node = OxmlElement("m:d")
    props = OxmlElement("m:dPr")
    beg = OxmlElement("m:begChr")
    beg.set(qn("m:val"), left)
    end = OxmlElement("m:endChr")
    end.set(qn("m:val"), right)
    props.extend([beg, end])
    e = OxmlElement("m:e")
    for item in m_expr(expression):
        e.append(item)
    node.extend([props, e])
    return node


def m_accent(expression, char: str = "̂"):
    node = OxmlElement("m:acc")
    props = OxmlElement("m:accPr")
    chr_node = OxmlElement("m:chr")
    chr_node.set(qn("m:val"), char)
    props.append(chr_node)
    e = OxmlElement("m:e")
    for item in m_expr(expression):
        e.append(item)
    node.extend([props, e])
    return node


def m_sum(index, expression, upper: str | None = None):
    node = OxmlElement("m:nary")
    props = OxmlElement("m:naryPr")
    char = OxmlElement("m:chr")
    char.set(qn("m:val"), "∑")
    loc = OxmlElement("m:limLoc")
    loc.set(qn("m:val"), "undOvr")
    props.extend([char, loc])
    if upper is None:
        sup_hide = OxmlElement("m:supHide")
        sup_hide.set(qn("m:val"), "1")
        props.append(sup_hide)
    sub = OxmlElement("m:sub")
    for item in m_expr(index):
        sub.append(item)
    sup = OxmlElement("m:sup")
    if upper is not None:
        for item in m_expr(upper):
            sup.append(item)
    e = OxmlElement("m:e")
    for item in m_expr(expression):
        e.append(item)
    node.extend([props, sub, sup, e])
    return node


def add_equation(container, expression, number: str) -> None:
    table = container.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    remove_table_borders(table)
    left, right = table.rows[0].cells
    left.width = Cm(15.2)
    right.width = Cm(1.1)
    set_cell_margins(left, 45, 45, 45, 45)
    set_cell_margins(right, 45, 45, 45, 45)
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    math = OxmlElement("m:oMath")
    for item in m_expr(expression):
        math.append(deepcopy(item))
    p._p.append(math)
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_after = Pt(0)
    run = rp.add_run(f"（{number}）")
    set_run_font(run, size=9)


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def set_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)


def set_portrait(section) -> None:
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.7)


def build_document() -> Path:
    OUT.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_sections(doc)
    configure_styles(doc)
    props = doc.core_properties
    props.title = "一种基于异构注视-动作数据逐样本监督路由的机器人操作策略训练方法及系统"
    props.subject = "Gaze-WAM 中国发明专利技术交底书"
    props.author = "Gaze-Project / Codex"
    props.keywords = "机器人操作, 注视, 异构数据, 监督路由, 扩散策略, 标签泄漏, K/V缓存"
    props.comments = f"代码基线 {CODE_COMMIT}，检索截止 {SEARCH_CUTOFF}"

    # Cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    add_top_rule(p, color=GREEN, size=34)
    r = p.add_run("GAZE-WAM  /  GAZE-PROJECT")
    set_run_font(r, size=10, bold=True)
    r.font.color.rgb = RGBColor.from_string(GREEN)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("中国发明专利技术交底书")
    set_run_font(r, size=15, bold=True)
    r.font.color.rgb = RGBColor.from_string(MUTED)

    title = doc.add_paragraph(style="Title")
    title.paragraph_format.space_before = Pt(6)
    title.paragraph_format.space_after = Pt(12)
    title.add_run("一种基于异构注视-动作数据\n逐样本监督路由的机器人操作策略训练方法及系统")

    deck = doc.add_paragraph(style="Deck")
    deck.add_run("代码对齐版  |  发明点筛选、现有技术检索、权利要求建议与实施证据")

    callout(
        doc,
        "申请结论",
        "建议申报主案。独立权利要求必须锁定：三类样本中注视标签在“输入条件”和“监督目标”之间的逐样本角色切换、可学习 [MASK] 条件、防标签泄漏互斥约束，以及按样本掩码聚合损失。不要把“异构缺标签多任务训练”或“共享 K/V 缓存”单独作为核心发明。",
        GREEN_LIGHT,
    )

    meta = [
        ("项目", "Gaze-WAM / Gaze-Project"),
        ("代码仓库", "thirdparty/gaze-dp"),
        ("代码分支", "gaze-wam-cleanup"),
        ("代码基线", CODE_COMMIT),
        ("检索截止", SEARCH_CUTOFF),
        ("文档状态", "内部评审稿；提交代理师前需补充申请人/发明人信息"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color=GRAY_MID, size=5)
    for i, (key, value) in enumerate(meta):
        table.cell(i, 0).width = Cm(3.5)
        table.cell(i, 1).width = Cm(12.8)
        set_cell_shading(table.cell(i, 0), GRAY_LIGHT)
        for cell in table.rows[i].cells:
            set_cell_margins(cell, 100, 140, 100, 140)
        p0 = table.cell(i, 0).paragraphs[0]
        p0.add_run(key)
        set_run_font(p0.runs[0], size=9, bold=True)
        p1 = table.cell(i, 1).paragraphs[0]
        p1.add_run(value)
        set_run_font(p1.runs[0], size=9)

    doc.add_paragraph()
    note = doc.add_paragraph(style="Small Gaze")
    note.add_run("保密提示：在首次专利申请日前，避免公开完整训练路由、互斥监督规则及代码实现细节。")
    page_break(doc)

    # Executive strategy
    heading(doc, "0. 发明点筛选与申请策略", 1)
    paragraph(
        doc,
        "本交底书不是项目功能清单，而是基于代码、实验记录、专利文献和论文公开状态筛选可形成稳定权利要求的技术组合。筛选标准包括：与现有技术的结构差异、差异解决的具体技术问题、代码可实施性、证据成熟度、规避设计难度和被无效风险。",
    )
    add_table(
        doc,
        ["候选发明点", "建议", "专利价值", "主要风险", "保护位置"],
        [
            ["注视-动作异构样本的三态逐样本监督路由与防泄漏", "主案", "高", "Google 缺标签多任务专利构成强背景技术；必须写入角色切换与互斥约束", "独立权利要求"],
            ["共享 world K/V cache 的动作/热图双扩散流", "纳入组合", "中", "Fast-WAM、AHA-WAM 已公开相近缓存与训练/推理解耦", "从属权利要求"],
            ["Frozen Cosmos 编解码的 256×16 热图潜变量", "实施例", "中低", "预训练图像 tokenizer 用于潜空间扩散较常见", "从属或说明书"],
            ["Collector 固定期限多速率融合与质量审计", "另案候选", "中", "多源机械臂采集、异步时间戳和周期误差校正已有直接专利；只能收窄到软件固定期限与联合质量门控", "单独交底书"],
            ["Quest-机器人坐标标定与 gaze 投影", "不优先", "中低", "3D gaze 控制机器人和 AR 坐标注册均已有强专利，当前差异主要是工程组合", "保留技术秘密"],
        ],
        widths_cm=[6.0, 1.8, 1.8, 5.5, 2.2],
        font_size=8.2,
    )
    heading(doc, "0.1 推荐保护层次", 2)
    numbered_list(
        doc,
        [
            "第一层：样本级状态路由。以数据源、动作标签、热图标签、注视标签和条件丢弃状态生成显式布尔门控。",
            "第二层：注视标签角色互斥。真实注视被用作动作条件时，禁止同一标签同时成为热图监督目标；条件被替换为 [MASK] 时，允许该标签转为热图监督。",
            "第三层：同一混合批次中的动作损失与热图损失分别按样本门控求 masked mean，并共同更新共享视觉/世界表征。",
            "第四层：共享世界 K/V 缓存、动作与热图目标流隔离、正式动作推理省略热图流，作为组合性从属特征。",
            "第五层：Frozen Cosmos 热图潜变量、DSNT 与空间 Jensen-Shannon 损失、具体维度和扩散步数作为实施例，不写死在最宽权利要求中。",
        ],
    )
    callout(
        doc,
        "最重要的撰写边界",
        "不得把技术方案概括为“将缺少不同标签的多个数据集联合训练共享网络”。US11717959B2 已公开非常接近的缺标签损失路由。可争辩的创造性来自注视标签的双重语义及其受控角色切换、[MASK] 条件和防泄漏契约的组合。",
        ORANGE_LIGHT,
    )

    heading(doc, "0.2 证据成熟度", 2)
    add_table(
        doc,
        ["证据项", "当前状态", "可支持的陈述", "不可支持的陈述"],
        [
            ["代码实现", "已实现并有显式运行时校验", "可实施的路由、损失和双流结构", "不等于已证明任务成功率提升"],
            ["模型训练", "HOT3D open-only 路径已端到端 smoke-trained", "开放注视数据链路可运行", "混合机器人数据训练已取得效果"],
            ["机器人/采集链路", "2026-07-22 真机记录 26/26 审计通过", "数据采集速率、丢包、写入延迟和有界遥操链路稳定", "主发明的模型学习效果已由真机实验验证"],
            ["聚焦契约测试", "2026-07-27：8 项中 6 通过；2 项失败来自 optional-presence mask 的旧预期与当前契约漂移", "[MASK]、distributed masked mean、world cache、open-only 与显式路由路径受聚焦测试覆盖", "完整历史测试套件全通过"],
            ["全量历史单测", "2026-07-27：197 passed、79 failed；失败集中在已移除的 Hydra 配置、旧 joint-transformer 契约、Cosmos JIT 夹具和 Windows 临时路径假设", "存在可定位的历史回归债务，不能据此否认当前代码中的路由结构", "生产训练链路或整个研究仓库已完成回归验证"],
        ],
        widths_cm=[3.0, 4.2, 5.1, 4.3],
        font_size=8.2,
    )

    # Main disclosure
    heading(doc, "1. 技术领域", 1)
    paragraph(
        doc,
        "本发明涉及机器人学习、计算机视觉、人机协同和多模态机器学习领域，具体涉及一种将无机器人动作标签的第一视角注视数据与包含机器人动作的示范数据组成异构混合批次，并依据逐样本标签可用性和注视条件状态，对动作监督与注视热图监督进行显式路由的机器人操作策略训练方法、训练系统、推理装置及计算机可读存储介质。",
    )

    heading(doc, "2. 背景技术", 1)
    paragraph(
        doc,
        "机器人操作策略通常依赖同步采集的图像、机器人状态和动作轨迹。此类示范数据成本高、场景覆盖有限。相对而言，公开第一视角数据集能够提供大量图像和人类注视标注，却不包含与目标机器人一致的动作标签。直接混合两类数据会遇到四个工程问题：缺失动作不能作为真实零动作参与监督；注视标签既可能作为策略条件，也可能作为注视预测目标；同一标签同时出现在输入和目标中会形成标签泄漏；辅助热图生成路径若进入正式动作推理，会引入不必要的实时开销。",
    )
    paragraph(
        doc,
        "已有多任务学习、缺标签学习和世界动作模型能够在不同任务损失之间共享网络，但未必规定注视标签在不同样本状态中的角色切换，也未必提供可执行的逐样本互斥契约。因此，需要一种能够在同一批次中可靠区分“注视是条件”与“注视是目标”的训练机制，并使开放数据和机器人数据对共享世界表征产生互补更新。",
    )

    heading(doc, "3. 现有技术检索", 1)
    heading(doc, "3.1 检索范围与关键词", 2)
    paragraph(
        doc,
        f"检索日期及截止日期均为 {SEARCH_CUTOFF}。检索覆盖 Google Patents 的专利全文、权利要求和专利族页面，以及 arXiv 与项目内研究记录。实际使用的核心检索式包括：robot gaze imitation learning policy；gaze saliency robot training；heterogeneous datasets missing labels / null label multi-task robot；robot gaze conditioned imitation learning policy；robot multi-rate sensor fusion camera timestamps；sensor data acquisition quality monitoring dropped frames queue latency timestamp；head-mounted gaze robot coordinate calibration gaze ray transformation；以及对应中文组合。法律状态以 Google Patents 页面在检索日的显示为线索，不替代专利局登记簿核验。",
    )
    paragraph(
        doc,
        "核验方法为：先按相关度筛选专利族，再打开具体公开文本，记录标题、申请人、最早优先权日、检索日法律状态，并逐项阅读独立权利要求；当关键缺标签机制仅出现在说明书而未进入独立权利要求时，单独标注说明书段落，避免把说明书公开误写为已授权权利要求范围。",
    )
    heading(doc, "3.2 最接近的专利文献", 2)

    add_table(
        doc,
        ["编号", "公开文本与时间", "已公开内容", "与本发明的差异", "风险"],
        [
            ["P1", "US20150339589A1\n优先权 2014-05-21；Brain Corp；Abandoned", "独立项1：由训练者 gaze 确定显著区域并更新学习参数，使机器人可脱离训练者执行任务", "未见开放无动作数据与机器人示范的三态角色路由、条件/目标互斥和逐样本损失门控", "中"],
            ["P2", "CN107097227B\n优先权 2017-04-17；原申请人北航；Active", "独立项1：追踪 gaze、识别关注目标、预测意图置信度并据此规划机器人运动；含 3D 环境映射", "偏向在线意图识别和规划，不是异构训练数据的逐样本监督路由", "低至中"],
            ["P3", "US11717959B2 / WO2019006091\n优先权 2017-06-28；Google；Active", "独立项1公开共享 joint network 上的抓取/语义损失；说明书[0008]-[0009]进一步公开缺标签或 null 标签样本只计算相应任务损失", "未公开 gaze 同时具有条件/目标双重语义、同标签互斥、条件 dropout 生成第三状态及对应防泄漏契约", "高"],
            ["P4", "JP2022115640A / JP7584134B2\n优先权 2021-01-28；东京大学；Active", "独立项1：按 gaze/peripheral 动作分类，分别生成高分辨率注视图和低分辨率周边图并训练两个模仿模型", "同一有动作示范按视野拆分，不是开放无动作 gaze 数据与机器人数据的混合监督", "中"],
            ["P5", "US12528186B2 / CN115551681B\n优先权 2020-05-14；Google；Active", "独立项12：从图像目标和自然语言目标两种机器人示范分别求损失，在共享潜在目标空间更新条件策略", "两个数据集均含动作轨迹；未见 gaze 条件/目标互斥和三态样本路由", "中高"],
        ],
        widths_cm=[1.0, 4.1, 4.4, 5.9, 1.2],
        font_size=7.7,
    )
    callout(
        doc,
        "检索结论",
        "P3 已足以否定把“异构缺标签样本 + 对应任务损失 + 共享网络”作为核心创新的稳妥性，P5 又强化了多数据集共享条件空间的背景。只有把三类样本状态、真实 gaze 输入与同标签热图目标互斥、可学习 [MASK] 替换、条件 dropout 和逐样本损失门控写成不可分割的组合，才保留较明确的新颖性与创造性争辩空间；这仍不是授权保证。",
        BLUE_LIGHT,
    )

    heading(doc, "3.3 关键论文与公开风险", 2)
    bullet_list(
        doc,
        [
            "Diffusion Policy（arXiv:2303.04137）公开动作扩散策略；因此“用扩散生成动作”本身不新。",
            "MimicPlay（arXiv:2302.12422）和 EgoMimic（arXiv:2410.24221）公开从人类/第一视角视频扩展机器人学习的方向；因此“利用开放第一视角数据”本身不新。",
            "gaze dual-resolution imitation learning（arXiv:2102.01295）、GazeBot（arXiv:2502.18121）及多任务真实机器人 gaze 数据工作（arXiv:2401.07603）说明 gaze 辅助机器人学习已有持续研究。",
            "Fast-WAM（arXiv:2603.16666）和 AHA-WAM（arXiv:2606.09811）对世界/动作双流、训练期世界分支和可复用 K/V 上下文构成强公开技术；共享 K/V cache 不宜单独主张。",
        ],
    )

    heading(doc, "3.4 全项目其他候选的权利要求级对比", 2)
    paragraph(
        doc,
        "为避免只在模型方案内部择优，另对 Collector 多速率采集/质量审计，以及 Quest-机器人坐标标定/gaze 投影进行了相同口径的专利检索。下列文献并不当然覆盖项目实现，但显著压缩了这些候选可主张的宽度。",
    )
    add_table(
        doc,
        ["编号", "文献与状态", "独立权利要求或关键公开", "对候选发明点的影响", "判断"],
        [
            ["P6", "CN109729278B\n优先权 2018-11-19；Momenta；Active", "独立项1：不同速率通道接收图像、其他传感器数据及时间戳，并做帧-时间戳映射和曝光时刻校正", "覆盖异步多传感器、不同传输速率和时间戳对应的宽泛框架", "Collector 强风险"],
            ["P7", "US10250868B1\n优先权 2017-03-29；Amazon；Active", "独立项1/5/14：由名义频率计算期望时间戳，与实测时间戳求误差并调整触发频率；说明书公开多传感器时间戳队列", "覆盖周期时间戳误差、传感器对齐和队列机制；但未见本项目的整套软件质量门控", "Collector 强风险"],
            ["P8", "CN121340281A\n优先权 2025-11-26；南昌大学等；Pending", "独立项1为多源机械臂操作采集；权利要求2及说明书公开 30 Hz 多相机、120 Hz 动捕、1 kHz 力传感、硬件同步和融合输出", "直接证明多速率机械臂示范采集本身高度拥挤；本项目只能聚焦 fixed-deadline 与质量闭环", "Collector 强风险"],
            ["P9", "US10157313B1\n优先权 2014-09-19；Colorado School of Mines；Active", "独立项1：跟踪眼动并转换为 3D gaze 控制输出；说明书用于机器人导航和物体操作", "3D gaze 到机器人控制和坐标表达已公开", "标定/投影强风险"],
            ["P10", "US11227441B2\n优先权 2019-07-04；Scopis/Stryker；Active", "独立项1：依据 AR 设备坐标系、参考坐标系、参考物体和用户观看方向校准坐标变换", "基于头戴设备观看方向校准跨坐标系注册已有直接方案", "标定/投影强风险"],
        ],
        widths_cm=[0.9, 3.8, 5.2, 5.1, 1.6],
        font_size=7.2,
    )
    callout(
        doc,
        "跨候选排序",
        "综合代码成熟度、可验证性、可规避难度和现有技术密度，Gaze-WAM 的防泄漏三态监督路由仍排名第一；Collector 固定期限质量闭环排名第二，适合另案但必须窄写；坐标标定、断电停录、一般安全包络和录像可视化属于更适合工程实现或技术秘密保护的候选，不建议挤入本主案。",
        GREEN_LIGHT,
    )

    heading(doc, "4. 要解决的技术问题", 1)
    numbered_list(
        doc,
        [
            "在同一训练批次内接纳标签集合不同的机器人示范数据和开放第一视角注视数据，同时避免把缺失动作解释为零动作。",
            "使注视标签在需要时作为动作策略条件，在另一些样本中作为热图监督目标，并阻止同一真实注视标签同时扮演两种角色。",
            "让动作监督和热图监督只作用于具有对应有效标签的样本，并以稳定的张量契约支持分布式训练。",
            "让热图辅助任务能够改善共享世界表征，同时不强制正式动作推理生成热图。",
            "对缺失字段、非法路由和零占位进行运行时校验，使错误在训练早期被确定性发现。",
        ],
    )

    heading(doc, "5. 术语与符号", 1)
    add_table(
        doc,
        ["符号/字段", "含义", "代码类型/约束"],
        [
            ["sᵢ / is_open", "样本是否来自开放第一视角数据", "[B] BoolTensor"],
            ["aᵢ / has_action", "动作标签是否可用", "[B] BoolTensor；开放行必须为假"],
            ["hᵢ / has_heatmap", "热图监督是否启用", "[B] BoolTensor"],
            ["gᵢ / has_gaze_label", "真实注视标签是否有效", "[B] BoolTensor"],
            ["cᵢ / use_gaze_condition", "是否把真实注视编码为输入条件", "为真时 gᵢ 必须为真"],
            ["dᵢ / is_gaze_condition_dropped", "真实注视条件是否被替换为 [MASK]", "当前实现要求 dᵢ = NOT cᵢ"],
            ["qᵢ", "归一化注视坐标", "[B,2]，有效值位于 [0,1]²"],
            ["z_mask", "可学习的缺失注视条件 token", "[1,1,D] 参数，经 batch 扩展"],
        ],
        widths_cm=[4.0, 7.4, 5.1],
        font_size=8.4,
    )

    heading(doc, "6. 技术方案", 1)
    heading(doc, "6.1 总体结构", 2)
    paragraph(
        doc,
        "训练系统读取机器人示范批次和开放注视批次，为每个样本生成显式路由标记；依据路由状态选择真实注视编码或可学习 [MASK] 条件；将图像 token 与注视条件 token 编码成共享世界表征；动作目标流和热图目标流分别读取该世界表征；最后按样本门控分别聚合动作损失和热图损失并联合更新。",
    )
    add_figure(doc, "figure_01_overall_training.png", "图1  异构注视-动作数据逐样本监督路由的总体训练架构")

    heading(doc, "6.2 三类样本状态", 2)
    add_table(
        doc,
        ["状态", "数据与条件", "动作监督", "热图监督", "技术含义"],
        [
            ["O：开放注视行", "无机器人动作；[MASK] 作条件", "关闭", "开启", "开放 gaze 标签只作为目标，不可泄漏到输入"],
            ["R0：机器人无注视行", "有动作；无有效 gaze；[MASK] 作条件", "开启", "关闭", "保留无 gaze 的机器人示范"],
            ["R1：机器人真实注视行", "真实 gaze 作动作条件", "开启", "关闭", "禁止复制输入 gaze 获得虚假热图低损失"],
            ["R2：机器人注视丢弃行", "有动作和 gaze；随机用 [MASK] 替代条件", "开启", "开启", "同一机器人样本同时提供动作和 gaze 目标，但 gaze 不进入条件"],
        ],
        widths_cm=[3.0, 4.7, 2.0, 2.0, 4.8],
        font_size=8.1,
    )
    paragraph(doc, "样本级门控定义如下，其中 aᵢ、hᵢ、gᵢ 均为标签可用性指示量：")
    add_equation(doc, [m_subsup("m", "i", "a"), "=", m_delim(["1−", m_sub("s", "i")]), m_sub("a", "i")], "1")
    add_equation(doc, [m_subsup("m", "i", "h"), "=", m_sub("h", "i"), m_sub("g", "i")], "2")
    add_equation(
        doc,
        [m_sub("c", "i"), "=", m_delim(["1−", m_sub("s", "i")]), m_sub("g", "i"), m_delim(["1−", m_sub("d", "i")])],
        "3",
    )
    paragraph(doc, "防标签泄漏约束可写为：任一以真实注视作为输入条件的样本，其热图监督门控必须为零。")
    add_equation(doc, [m_sub("c", "i"), m_subsup("m", "i", "h"), "=0"], "4")
    add_figure(doc, "figure_02_sample_routing.png", "图2  单一样本的监督状态判定与防标签泄漏流程")

    heading(doc, "6.3 注视条件编码", 2)
    paragraph(
        doc,
        "有效注视坐标 qᵢ 经空间基函数编码器和投影网络得到注视 token。若路由规定不使用真实注视条件，则用同一可训练 [MASK] token 替代。编码器还校验 use_gaze_condition 不得在 has_gaze_label=False 的行上为真。",
    )
    add_equation(
        doc,
        [m_subsup("z", "i", "g"), "=", m_sub("c", "i"), m_sub("E", "g"), m_delim(m_sub("q", "i")), "+", m_delim(["1−", m_sub("c", "i")]), m_sub("z", "mask")],
        "5",
    )

    heading(doc, "6.4 双目标扩散与共享 timestep", 2)
    paragraph(
        doc,
        "动作目标和热图潜变量目标在训练时使用同一批次 timestep，但分别采样噪声。令 r∈{a,h} 表示动作流或热图流，其前向扩散可表示为：",
    )
    add_equation(
        doc,
        [m_subsup("x", "t", "r"), "=", m_sqrt(m_sub("ᾱ", "t")), m_subsup("x", "0", "r"), "+", m_sqrt(["1−", m_sub("ᾱ", "t")]), m_sup("ε", "r")],
        "6",
    )

    heading(doc, "6.5 共享世界 K/V 缓存与目标流隔离", 2)
    paragraph(
        doc,
        "图像 token 和注视/[MASK] token 组成稳定条件序列 W。七层 ContextSelfBlock 逐层计算世界表示，并导出每层 K/V。动作流和热图流分别用自身 target token 产生查询，同时拼接对应层 world K/V 与本流自身 K/V。两个目标流不读取对方 target token。",
    )
    add_equation(
        doc,
        [m_delim([m_subsup("K", "l", "w"), ",", m_subsup("V", "l", "w")]), "=", m_sub("KV", "l"), m_delim(m_sub("W", "l"))],
        "7",
    )
    add_equation(
        doc,
        [
            m_subsup("Y", "l", "r"),
            "=softmax",
            m_delim(
                m_frac(
                    [m_subsup("Q", "l", "r"), m_sup(m_delim([m_subsup("K", "l", "w"), ";", m_subsup("K", "l", "r")], "[", "]"), "T")],
                    m_sqrt("d"),
                )
            ),
            m_delim([m_subsup("V", "l", "w"), ";", m_subsup("V", "l", "r")], "[", "]"),
        ],
        "8",
    )
    add_figure(doc, "figure_03_dual_stream_cache.png", "图3  共享世界 K/V 缓存的动作-注视热图双流扩散网络")

    heading(doc, "6.6 按样本 masked mean 损失", 2)
    paragraph(
        doc,
        "代码先把每个样本的动作预测误差或热图误差在非 batch 维度求均值，再依据门控执行 distributed_masked_mean。由此，零占位仅保持批次张量同形状，不会进入对应损失。",
    )
    add_equation(
        doc,
        [m_sub("L", "a"), "=", m_frac(m_sum("i", [m_subsup("m", "i", "a"), m_subsup("ℓ", "i", "a")]), [m_sum("i", m_subsup("m", "i", "a")), "+δ"])],
        "9",
    )
    add_equation(
        doc,
        [m_sub("L", "h"), "=", m_frac(m_sum("i", [m_subsup("m", "i", "h"), m_subsup("ℓ", "i", "h")]), [m_sum("i", m_subsup("m", "i", "h")), "+δ"])],
        "10",
    )
    paragraph(
        doc,
        "当前主配置采用 Frozen Cosmos 解码后的热图图像，通过 intensity_softplus 构造空间分布。φ 表示零校准 softplus 与非负归一化组合：",
    )
    add_equation(
        doc,
        [
            m_sub("P", "i"),
            m_delim("u,v"),
            "=",
            m_frac(
                ["φ", m_delim([m_sub("H", "i"), m_delim("u,v")])],
                [m_sum("u′,v′", ["φ", m_delim([m_sub("H", "i"), m_delim("u′,v′")])]), "+δ"],
            ),
        ],
        "11",
    )
    add_equation(
        doc,
        [
            m_sub(m_accent("q"), "i"),
            "=",
            m_sum(
                "u,v",
                [
                    m_sub("P", "i"),
                    m_delim("u,v"),
                    m_delim([m_frac("u+1/2", "W"), ",", m_frac("v+1/2", "H")], "[", "]"),
                ],
            ),
        ],
        "12",
    )
    add_equation(
        doc,
        [m_sub("L", "h"), "=", m_sub("λ", "xy"), m_sub("L", "xy"), "+", m_sub("λ", "JS"), m_sub("L", "JS")],
        "13",
    )
    add_equation(
        doc,
        ["L=", m_sub("λ", "a"), m_sub("L", "a"), "+", m_sub("λ", "h"), m_delim([m_sub("λ", "xy"), m_sub("L", "xy"), "+", m_sub("λ", "JS"), m_sub("L", "JS")])],
        "14",
    )

    heading(doc, "6.7 训练路径与正式动作推理路径", 2)
    paragraph(
        doc,
        "训练时两个目标流均存在，以各自门控产生梯度。正式动作推理时，系统仅预填充一次 world cache，并在八步 DDIM 去噪中重复读取该缓存；调用约束要求 noisy_heatmap=None，热图解码器和热图输出完全省略。",
    )
    add_equation(doc, [m_sub("x", "heatmap"), "=∅⇒", m_sub("Y", "heatmap"), "=∅"], "15")
    add_figure(doc, "figure_04_train_inference.png", "图4  联合训练路径与正式动作推理路径")

    heading(doc, "7. 有益效果", 1)
    numbered_list(
        doc,
        [
            "扩大可用训练数据范围：无机器人动作的开放第一视角 gaze 样本能够更新热图目标流和共享视觉/世界表征，而不会污染动作目标。",
            "避免条件复制：真实 gaze 作为动作条件的行禁止热图监督，阻止模型仅复制输入标签获得低热图损失。",
            "提升缺失条件鲁棒性：随机将机器人真实 gaze 替换为可学习 [MASK]，使同一动作策略能够在有 gaze 和无 gaze 条件下训练。",
            "保持梯度语义清晰：每个损失先按样本计算，再以显式布尔 mask 聚合；缺失标签的零占位不会被误当成监督。",
            "隔离目标流：动作 target token 与热图 target token 不交叉读取，开放数据中的 dummy action 不会泄漏到热图学习。",
            "降低正式推理负担：稳定世界条件只预填充一次；热图辅助流在正式动作推理中跳过。",
            "提高工程可审计性：批次字段的 dtype、shape、零占位和互斥关系均在运行时校验。",
        ],
    )
    callout(
        doc,
        "效果证据边界",
        "上述第1至第7项中，路由、隔离和推理省略由代码结构直接支持；“提高任务成功率”“提高泛化精度”等模型效果尚需混合机器人数据训练和对照实验验证，现阶段只能作为预期技术效果，不应写成已实测结论。",
        ORANGE_LIGHT,
    )

    heading(doc, "8. 具体实施方式", 1)
    heading(doc, "8.1 训练流程", 2)
    numbered_list(
        doc,
        [
            "S101：读取机器人示范样本。每个样本至少包含图像和动作块，可选包含归一化 gaze 坐标和热图。",
            "S102：读取开放第一视角样本。每个样本包含图像和 gaze 标签，不包含机器人动作；为保持张量同形状，构造零动作占位。",
            "S103：依据数据源和字段存在性生成 is_open、has_action、has_heatmap、has_gaze_label。",
            "S104：对具有 gaze 标签的机器人样本，以预设概率采样 gaze-condition dropout；默认概率为 0.2。",
            "S105：生成 use_gaze_condition 和 is_gaze_condition_dropped，并执行批次契约校验。",
            "S106：DINOv3 图像编码器把两帧观测分别转换为 256 个 patch token；注视编码器输出真实 gaze token 或 learned [MASK] token。",
            "S107：共享世界塔预填充七层 K/V；动作和热图目标流用同一 timestep、独立噪声进行扩散训练。",
            "S108：动作误差仅在 (~is_open)&has_action 的样本上聚合；热图误差仅在 has_heatmap&has_gaze_label 的样本上聚合。",
            "S109：按权重组合损失并反向传播，共同更新视觉编码器、注视编码器、世界塔和对应目标头。",
        ],
    )
    heading(doc, "8.2 与当前代码一致的路由伪代码", 2)
    add_code_block(
        doc,
        "for each sample i in mixed_batch:\n"
        "    if is_open[i]:\n"
        "        has_action[i] = False\n"
        "        use_gaze_condition[i] = False\n"
        "        has_heatmap[i] = True\n"
        "    elif not has_gaze_label[i]:\n"
        "        has_action[i] = True\n"
        "        use_gaze_condition[i] = False\n"
        "        has_heatmap[i] = False\n"
        "    elif gaze_condition_is_dropped[i]:\n"
        "        has_action[i] = True\n"
        "        use_gaze_condition[i] = False\n"
        "        has_heatmap[i] = True\n"
        "    else:\n"
        "        has_action[i] = True\n"
        "        use_gaze_condition[i] = True\n"
        "        has_heatmap[i] = False\n"
        "\n"
        "action_mask  = (~is_open) & has_action\n"
        "heatmap_mask = has_heatmap & has_gaze_label",
    )

    heading(doc, "8.3 主配置实施例", 2)
    add_table(
        doc,
        ["参数", "当前值", "代码位置", "权利要求处理"],
        [
            ["观测帧数", "2", "diffusion_policy/config/task/gaze_wam.yaml:9", "实施例，不写死"],
            ["动作块", "[16,10]", "diffusion_policy/config/task/gaze_wam.yaml:13-14", "实施例，不写死"],
            ["热图潜变量", "[256,16]", "task 配置与 Cosmos CI16x16", "从属/实施例"],
            ["Transformer", "7 层、8 heads、768 embedding", "diffusion_policy/config/train_gaze_wam_workspace.yaml:12,77-79", "从属/实施例"],
            ["训练扩散步", "50", "diffusion_policy/config/train_gaze_wam_workspace.yaml:20", "实施例"],
            ["推理步", "8", "diffusion_policy/config/train_gaze_wam_workspace.yaml:61", "实施例"],
            ["gaze dropout", "0.2", "diffusion_policy/config/task/gaze_wam.yaml:44", "可写概率区间"],
            ["热图目标", "DSNT XY + spatial JS", "diffusion_policy/config/train_gaze_wam_workspace.yaml:66-69；diffusion_policy/model/gaze_wam/loss.py", "从属"],
            ["热图 codec", "Frozen Cosmos CI16x16", "diffusion_policy/config/train_gaze_wam_workspace.yaml:89-101", "从属"],
            ["CFG scale", "1.0", "diffusion_policy/config/train_gaze_wam_workspace.yaml:75", "实施例"],
        ],
        widths_cm=[3.2, 3.0, 7.0, 3.2],
        font_size=8.1,
    )

    heading(doc, "8.4 运行时契约", 2)
    bullet_list(
        doc,
        [
            "所有路由标记必须是 batch 维 [B] BoolTensor。",
            "开放行 has_action 必须为假，机器人行 has_action 必须为真。",
            "开放行 use_gaze_condition 必须为假，且必须启用热图目标。",
            "use_gaze_condition 为真的行必须具有有效 gaze 标签。",
            "机器人真实 gaze 条件行禁止 has_heatmap 为真。",
            "has_action=False 的动作行和 has_heatmap=False 的热图行必须是同形状零占位。",
            "无 gaze 标签行的 gaze_xy 必须是零占位；is_gaze_condition_dropped 必须等于 use_gaze_condition 的逻辑非。",
        ],
    )

    heading(doc, "9. 可替代实施例", 1)
    bullet_list(
        doc,
        [
            "图像编码器可替换为其他视觉 Transformer、卷积网络或视觉语言编码器，只要输出可供共享世界塔使用的 token。",
            "注视标签可为二维点、多个 fixation、概率密度、注视轨迹或对象级注意目标；条件编码与目标生成应保持角色互斥。",
            "[MASK] 可以是单个可学习 token、按场景类别选择的 token、从先验分布采样的 token，或由缺失条件编码器生成的 token。",
            "动作目标可为关节位置、关节速度、末端位姿、相对位移、夹爪状态或其组合；只要动作损失由逐样本动作可用性门控。",
            "热图目标可在像素空间、潜变量空间或离散 token 空间扩散；损失可采用 MSE、KL、JS、点 NLL、DSNT 或其组合。",
            "世界表征可通过完整注意力、交叉注意力、K/V cache、状态空间模型或其他条件记忆实现；共享 K/V 是优选实施例。",
            "gaze-condition dropout 可按固定概率、课程策略、样本置信度或任务阶段确定。",
            "混合批次可在单机、数据并行或模型并行环境中构造；masked mean 的分子和分母可跨设备归约。",
        ],
    )

    heading(doc, "10. 权利要求书建议", 1)
    callout(
        doc,
        "撰写提示",
        "以下为技术性权利要求骨架，尚需由专利代理师根据申请主体、目标法域和检索报告调整法律用语。独立权利要求不应删除三类路由和互斥约束，否则容易落入 P3/P5 的宽泛缺标签多任务框架。",
        BLUE_LIGHT,
    )
    heading(doc, "10.1 方法独立权利要求建议", 2)
    claims = [
        "1. 一种机器人操作策略训练方法，包括：获取第一样本集和第二样本集，第一样本集的样本包括图像、机器人动作标签以及可选注视标签，第二样本集的样本包括图像和注视标签且不包括机器人动作标签；将所述样本组成混合批次，并针对每个样本生成数据源标记、动作标签可用性标记、注视标签可用性标记、热图监督标记和注视条件使用标记；依据所述标记将样本至少路由为开放注视样本、机器人真实注视条件样本和机器人注视条件丢弃样本，其中开放注视样本以掩码注视 token 作为条件、关闭动作监督并开启热图监督，机器人真实注视条件样本以真实注视标签作为条件、开启动作监督并关闭由同一注视标签形成的热图监督，机器人注视条件丢弃样本以所述掩码注视 token 替代真实注视条件并同时开启动作监督和热图监督；由图像和所选择的注视条件生成共享世界表征；基于所述共享世界表征生成动作预测和注视热图预测；分别依据动作标签可用性和热图监督标记对逐样本动作损失和逐样本热图损失进行掩码聚合，并依据聚合损失更新机器人操作策略。",
        "2. 根据权利要求1所述的方法，其中，机器人样本不具有有效注视标签时，以所述掩码注视 token 作为条件，开启动作监督并关闭热图监督。",
        "3. 根据权利要求1所述的方法，其中，注视条件使用标记与热图监督标记满足互斥约束，使同一真实注视标签用作输入条件时不被同时用作热图监督目标。",
        "4. 根据权利要求1所述的方法，其中，所述掩码注视 token 为随训练更新的可学习参数。",
        "5. 根据权利要求1所述的方法，其中，对具有有效注视标签的机器人样本按预设概率执行注视条件丢弃，以在保留动作监督的同时将对应注视标签转为热图监督目标。",
        "6. 根据权利要求1所述的方法，其中，共享世界表征由多个上下文层生成，每一上下文层输出世界键值缓存，动作目标流和热图目标流分别以自身目标 token 为查询读取所述世界键值缓存。",
        "7. 根据权利要求6所述的方法，其中，动作目标流不读取热图目标 token，热图目标流不读取动作目标 token。",
        "8. 根据权利要求6所述的方法，其中，在正式动作推理时预填充一次世界键值缓存，并在多个动作去噪步骤中复用该缓存，且省略热图目标流。",
        "9. 根据权利要求1所述的方法，其中，注视热图经冻结的图像编解码器转换为潜变量 token，并在潜变量空间执行扩散训练。",
        "10. 根据权利要求1所述的方法，其中，热图损失包括由空间概率分布期望获得的注视坐标损失和预测热图分布与目标热图分布之间的 Jensen-Shannon 损失。",
        "11. 根据权利要求1所述的方法，其中，缺失动作、缺失热图和缺失注视坐标使用与有效数据同形状的零占位，并通过独立布尔标记排除于相应损失之外。",
    ]
    for claim in claims:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.first_line_indent = Cm(-0.4)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_together = True
        r = p.add_run(claim)
        set_run_font(r, size=9.5)

    heading(doc, "10.2 系统与介质权利要求建议", 2)
    system_claims = [
        "12. 一种机器人操作策略训练系统，包括数据读取模块、逐样本路由模块、注视条件编码模块、共享世界表征模块、动作预测模块、注视热图预测模块和损失聚合模块，各模块被配置为执行权利要求1至11任一项所述的方法。",
        "13. 根据权利要求12所述的系统，其中，共享世界表征模块被配置为输出逐层世界键值缓存，动作预测模块和注视热图预测模块分别读取所述缓存且彼此不读取对方目标 token。",
        "14. 一种机器人操作装置，包括处理器和存储器，存储器中存储的程序被处理器执行时，使处理器在训练阶段执行权利要求1至11任一项所述的方法，并在推理阶段省略注视热图目标流而输出机器人动作。",
        "15. 一种计算机可读存储介质，其上存储有程序，所述程序被处理器执行时实现权利要求1至11任一项所述的方法。",
    ]
    for claim in system_claims:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.first_line_indent = Cm(-0.4)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_together = True
        r = p.add_run(claim)
        set_run_font(r, size=9.5)

    heading(doc, "11. 新颖性、创造性与授权风险分析", 1)
    heading(doc, "11.1 新颖性判断", 2)
    paragraph(
        doc,
        "基于当前检索，尚未发现单一文献同时公开以下全部必要特征：开放无动作 gaze 行使用 [MASK] 条件并仅监督热图；机器人真实 gaze 行使用真实 gaze 条件并仅监督动作；机器人 gaze dropout 行使用 [MASK] 条件并同时监督动作和热图；真实 gaze 输入与同标签热图目标互斥；两类损失按逐样本 mask 聚合。因而该组合具备可主张的新颖性空间。",
    )
    heading(doc, "11.2 创造性论证", 2)
    paragraph(
        doc,
        "最接近的 P3 在独立项1中以抓取损失和语义损失训练共享 joint network，并在说明书[0008]-[0009]明确处理缺标签或 null 标签样本；P5 独立项12又公开了对两个机器人示范数据集分别求损失并更新共享条件策略。因此，客观技术问题不能表述成一般的异构或缺标签联合训练，而应表述为：当同一 gaze 标签既可能是机器人动作策略的输入条件、又可能是辅助 gaze 预测目标时，如何联合利用开放无动作 gaze 数据和机器人示范，同时阻断条件到目标的标签复制，并避免 dummy action 进入动作监督。P3/P5 均未给出把同一 gaze 标签按样本状态在条件与目标之间切换、以条件 dropout 构造第三类监督状态并强制互斥的技术动机；P1/P2/P4 也未补足这一组合。因此，创造性论证应始终围绕标签双重语义冲突与可执行防泄漏契约，而不是一般缺标签学习。",
    )
    heading(doc, "11.3 主要授权风险", 2)
    add_table(
        doc,
        ["风险", "概率", "后果", "应对"],
        [
            ["P3/P5 被用于覆盖宽泛异构多任务训练", "高", "宽权利要求缺乏新颖性或创造性", "在独立项保留三态路由、角色互斥和 [MASK] 条件"],
            ["Fast-WAM/AHA-WAM 覆盖共享 K/V 和推理跳过世界分支", "高", "缓存特征不能单独支撑创造性", "仅作为从属组合，主案核心放在监督路由"],
            ["说明书缺少混合训练效果数据", "中", "效果表述受质疑", "把代码可实施效果与预测精度效果分开；补做消融"],
            ["术语过度绑定代码字段", "中", "保护范围易规避", "权利要求使用功能性标记，说明书给出字段名作为实施例"],
            ["过早公开 2026 年论文或仓库", "中", "影响部分法域新颖性", "尽快确定首次申请日并核验所有公开时间"],
        ],
        widths_cm=[5.0, 1.5, 4.5, 5.5],
        font_size=8.2,
    )

    section_12_heading = doc.add_heading("12. 现有技术区别矩阵", level=1)
    section_12_heading.paragraph_format.page_break_before = True
    paragraph(doc, "● 表示明确公开或高度相近；△ 表示部分相关；空白表示当前检索未见明确公开。矩阵用于技术筛选，不是法律上的逐项侵权或无效结论。")
    heading(doc, "12.1 专利文献对比", 2)
    add_table(
        doc,
        ["必要特征", "P1\nUS20150339589", "P2\nCN107097227", "P3\nUS11717959", "P4\nJP7584134", "P5\nUS12528186"],
        [
            ["开放第一视角 gaze、无机器人动作", "△", "", "△", "", "△"],
            ["机器人图像+动作+可选 gaze", "△", "●", "△", "●", "●"],
            ["逐样本显式 presence/source 标记", "", "", "●", "", "●"],
            ["开放行：[MASK] 条件、仅热图监督", "", "", "", "", ""],
            ["机器人真实 gaze：条件、仅动作监督", "△", "△", "", "△", ""],
            ["机器人 gaze dropout：动作+热图监督", "", "", "", "", ""],
            ["同一 gaze 输入/目标互斥防泄漏", "", "", "", "", ""],
            ["按样本 masked mean 联合更新", "", "", "●", "", "●"],
            ["共享 world K/V 双目标流", "", "", "", "", ""],
            ["正式动作推理省略热图流", "", "", "", "", ""],
        ],
        widths_cm=[6.0, 2.1, 2.1, 2.1, 2.1, 2.1],
        font_size=7.3,
    )
    page_break(doc)
    heading(doc, "12.2 世界模型论文与本方案对比", 2)
    add_table(
        doc,
        ["必要特征", "Fast-WAM / AHA-WAM", "本方案"],
        [
            ["开放第一视角 gaze、无机器人动作", "", "●"],
            ["机器人图像+动作+可选 gaze", "●", "●"],
            ["逐样本显式 presence/source 标记", "", "●"],
            ["开放行：[MASK] 条件、仅热图监督", "", "●"],
            ["机器人真实 gaze：条件、仅动作监督", "", "●"],
            ["机器人 gaze dropout：动作+热图监督", "", "●"],
            ["同一 gaze 输入/目标互斥防泄漏", "", "●"],
            ["按样本 masked mean 联合更新", "△", "●"],
            ["共享 world K/V 双目标流", "●", "●"],
            ["正式动作推理省略热图流", "●", "●"],
        ],
        widths_cm=[7.0, 4.8, 4.8],
        font_size=8.0,
    )

    heading(doc, "13. 代码与证据映射", 1)
    paragraph(doc, f"以下映射对应分支 gaze-wam-cleanup、提交 {CODE_COMMIT}。行号以该提交为准。")
    add_table(
        doc,
        ["技术特征", "代码位置", "关键行为"],
        [
            ["开放样本路由", "diffusion_policy/dataset/gaze_wam_mixing.py:182-203", "is_open=True；has_action=False；has_heatmap=True；use_gaze_condition=False"],
            ["机器人 gaze dropout", ".../dataset/gaze_wam_mixing.py:264-298", "基于概率与 gaze 可用性生成三类机器人状态"],
            ["混合批次标记", ".../dataset/gaze_wam_mixing.py:397-427", "拼接六个显式路由标记并生成零占位"],
            ["可学习 [MASK] 与校验", ".../model/gaze_wam/gaze_encoder.py:107-171", "真实 gaze/掩码 token 选择；禁止无标签行使用真实条件"],
            ["路由契约摘要", ".../policy/gaze_wam_policy.py:1277-1373", "列出三类主状态和损失 mask"],
            ["运行时防泄漏校验", ".../policy/gaze_wam_policy.py:1473-1539", "校验 shape、dtype、零占位、互斥和 dropped=~use"],
            ["按样本损失", ".../policy/gaze_wam_policy.py:1903-2177", "action mask 与 heatmap mask 分开 masked mean"],
            ["world cache 预填充", ".../model/gaze_wam/cached_dual_stream_transformer.py:410-450", "图像+gaze 条件逐层导出 K/V"],
            ["双目标流与推理省略", ".../model/gaze_wam/cached_dual_stream_transformer.py:532-663", "动作/热图不交叉；inference 禁止 noisy_heatmap"],
            ["Frozen Cosmos codec", ".../model/gaze_wam/heatmap_decoder.py:18-241", "校验 JIT 路径和潜变量 shape，冻结 tokenizer"],
            ["空间损失", ".../model/gaze_wam/loss.py:80-224", "intensity_softplus、DSNT、JS"],
        ],
        widths_cm=[4.0, 7.2, 5.2],
        font_size=7.9,
    )

    heading(doc, "14. 实验记录与可验证性", 1)
    heading(doc, "14.1 模型侧状态", 2)
    paragraph(
        doc,
        "README 在 2026-06-27 记录：HOT3D → zarr → train 的 open-only 路径可端到端运行，并已使用验证数据 smoke-train；该时点机器人数据收集尚未开始。此后项目已于 2026-07-22 获得并审计真实 Quest/机器人/RealSense 原始记录，但 Collector 正式记录目录到 Gaze-WAM canonical robot zarr 的转换桥仍未完成，尚未形成可用于本主案的混合机器人/action 训练结果。因此应表述为“机器人原始示范已采集、训练数据转换和混合训练尚未完成”，不得声称主模型已提升真实机器人任务成功率。",
    )
    heading(doc, "14.2 真机采集链路证据", 2)
    paragraph(
        doc,
        "记录 record_bounded_teleop_v3_20260722_180900 的 performance_audit.json 显示 26/26 检查通过。该记录证明采集基础设施能够生成稳定的机器人/视觉/Quest 对齐数据，可作为实施条件证据，但不是主发明的训练效果实验。",
    )
    add_table(
        doc,
        ["指标", "实测值", "结论"],
        [
            ["Flexiv robot state", "89.99995 Hz；p95 gap 11.5 ms；missed ticks 0", "通过 90 Hz 稳定性阈值"],
            ["最终融合样本", "29.99755 Hz；p95 lateness 8.9 ms；missed ticks 0", "通过 30 Hz 稳定性阈值"],
            ["RealSense end / third RGB", "29.97861 / 29.97905 Hz", "两路均稳定接近 30 Hz"],
            ["相机 queue drops", "0 / 0", "无写队列丢帧"],
            ["capture-to-write p95", "3.8447 / 4.2059 ms", "低于 100 ms 阈值"],
            ["Quest UDP", "3934 datagrams；loss/duplicate/reorder 均为 0", "传输通过"],
            ["Quest source age p95", "27.391 ms", "低于 60 ms 阈值"],
            ["TCP 有界遥操范围", "40.500 mm / 7.957°", "低于 50 mm / 10° 测试范围"],
            ["响应相关性", "最佳相关 lag 0.280 s；correlation 0.681", "仅表示目标-机器人响应相关滞后，不等于绝对网络时延"],
        ],
        widths_cm=[4.4, 7.0, 5.0],
        font_size=8.1,
    )
    paragraph(
        doc,
        "证据文件：artifacts/record_bounded_teleop_v3_20260722_180900/performance_audit.json、bounded_teleop_probe_report.json、teleop_latency_analysis.json 和 pc_session_summary.json。",
        style="Small Gaze",
    )

    heading(doc, "14.3 建议补充的专利实验", 2)
    add_table(
        doc,
        ["实验", "对照", "指标", "作用"],
        [
            ["三态路由消融", "无 R2；无互斥；无 [MASK]", "动作成功率、热图 DSNT/JS、gaze dependency", "证明主发明组合效果"],
            ["开放数据比例", "0%、25%、50%、75%", "机器人任务成功率、跨场景泛化", "证明开放 gaze 数据贡献"],
            ["标签泄漏检查", "R1 同时监督热图 vs 禁止", "训练/验证热图差距、遮蔽 gaze 后性能", "证明互斥约束必要性"],
            ["缓存推理消融", "每步重算 world vs cache；含热图流 vs 跳过", "单步延迟、吞吐、显存", "支撑从属缓存效果"],
            ["缺 gaze 鲁棒性", "真实 gaze、[MASK]、噪声 gaze", "动作成功率与输出变化", "证明 dropout 技术效果"],
        ],
        widths_cm=[3.5, 4.5, 4.3, 4.2],
        font_size=8.1,
    )

    heading(doc, "15. 分案与后续布局建议", 1)
    numbered_list(
        doc,
        [
            "主案：本交底书的三态逐样本监督路由、防 gaze 标签泄漏和动作/热图联合训练。",
            "分案候选 A：固定期限调度的 Flexiv 90 Hz、双 RealSense 30 Hz、Quest UDP 与 30 Hz 融合采样，以及 source age、missed tick、queue drop 和 capture-to-write 联合审计。该主题解决实时采集完整性问题，与模型训练主题不同；受 P6-P8 限制，独立项必须强调软件 fixed-deadline 时间线、异速率原始流保留、按目标时刻选择/标注来源样本和多指标联合准入，不能宽写成一般多传感器同步采集。",
            "分案候选 B：若后续形成稳定实验，可围绕“开放第一视角 gaze 预训练 + 机器人少样本适配”的具体阶段化训练和冻结/解冻策略另案。",
            "不建议单独申请：共享 K/V cache、Frozen Cosmos codec、普通 diffusion action policy、一般 Quest-机器人坐标标定、断电停录和矩形安全包络；这些方向已有强论文、专利或通用工程做法。",
            "申请节奏：先核验所有代码仓库、论文、演示和会议材料的公开日期；在新公开前提交中国首次申请，并在 12 个月优先权期内决定 PCT/海外布局。",
        ],
    )

    heading(doc, "16. 参考文献与链接", 1)
    refs = [
        ("[P1] US20150339589A1, Brain Corporation", "https://patents.google.com/patent/US20150339589A1/en"),
        ("[P2] CN107097227B", "https://patents.google.com/patent/CN107097227B/en"),
        ("[P3] US11717959B2 / WO2019006091 family", "https://patents.google.com/patent/US11717959B2/en"),
        ("[P4] JP2022115640A / JP7584134B2", "https://patents.google.com/patent/JP2022115640A/en"),
        ("[P5-CN] CN115551681B", "https://patents.google.com/patent/CN115551681B/en"),
        ("[P5-US] US12528186B2", "https://patents.google.com/patent/US12528186B2/en"),
        ("[P6] CN109729278B, rate-configurable multi-sensor receiving", "https://patents.google.com/patent/CN109729278B/en"),
        ("[P7] US10250868B1, synchronizing data streams", "https://patents.google.com/patent/US10250868B1/en"),
        ("[P8] CN121340281A, multi-source robotic-arm data acquisition", "https://patents.google.com/patent/CN121340281A/en"),
        ("[P9] US10157313B1, 3D gaze control of robot", "https://patents.google.com/patent/US10157313B1/en"),
        ("[P10] US11227441B2, AR registration calibration", "https://patents.google.com/patent/US11227441B2/en"),
        ("[L1] Diffusion Policy", "https://arxiv.org/abs/2303.04137"),
        ("[L2] MimicPlay", "https://arxiv.org/abs/2302.12422"),
        ("[L3] EgoMimic", "https://arxiv.org/abs/2410.24221"),
        ("[L4] Gaze dual-resolution imitation learning", "https://arxiv.org/abs/2102.01295"),
        ("[L5] GazeBot", "https://arxiv.org/abs/2502.18121"),
        ("[L6] Multi-task real-robot gaze dataset", "https://arxiv.org/abs/2401.07603"),
        ("[L7] Fast-WAM", "https://arxiv.org/abs/2603.16666"),
        ("[L8] AHA-WAM", "https://arxiv.org/abs/2606.09811"),
    ]
    for label, url in refs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        add_hyperlink(p, label, url)
        tail = p.add_run(f"  {url}")
        set_run_font(tail, size=8.5)
        tail.font.color.rgb = RGBColor.from_string(MUTED)

    heading(doc, "17. 法律与检索限制说明", 1)
    callout(
        doc,
        "重要说明",
        "本文件是基于公开数据库、论文和代码的技术分析，不构成法律意见，也不能替代 CNIPA、WIPO、USPTO、J-PlatPat 等权威数据库的专业查新、法律状态核验、权利要求解释或自由实施（FTO）分析。正式申请前应由专利代理师复核专利族、优先权、同族权利要求、审查历史、法律状态和申请主体信息。",
        RED_LIGHT,
    )

    heading(doc, "附录 A：交底完整性检查表", 1)
    add_table(
        doc,
        ["项目", "状态", "提交代理师前动作"],
        [
            ["发明名称与核心组合", "已确定", "保持三态路由 + 互斥 + [MASK] + masked loss"],
            ["代码对应", "已映射到 commit", "冻结一份申请日代码归档和哈希"],
            ["附图", "4 页 Draw.io + PNG", "代理师可直接调整标号，不改变数据流语义"],
            ["现有技术", "已完成公开数据库权利要求级复核", "由代理师在 CNIPA/商业数据库复核专利族、审查历史与非专利文献"],
            ["模型效果实验", "未完成混合机器人训练", "补三态消融、标签泄漏和缓存延迟实验"],
            ["发明人/申请人", "待填写", "按实质贡献确认，不以职位代替贡献"],
            ["首次公开日", "待统一核验", "核对 Git、论文、会议、视频和演示材料"],
        ],
        widths_cm=[4.2, 4.0, 8.2],
        font_size=8.3,
    )

    # Headers/footers must be applied after all sections exist.
    add_headers_footers(doc)
    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    print(build_document())
